#!/usr/bin/env python3
"""
Create reference.docx template for Vietnamese thesis format.
Requires: python-docx (pip install python-docx)

Vietnamese thesis specifications:
- Font: Times New Roman
- Body text: 13pt
- Headings: 14-18pt
- Left margin: 3.5 cm
- Right margin: 2.0 cm
- Top margin: 3.0 cm
- Bottom margin: 3.0 cm
- Line spacing: 1.5 lines
- Page size: A4 (210 × 297 mm)
"""

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed")
    print("Install with: pip install python-docx")
    exit(1)

def set_font(run, font_name='Times New Roman', size=13):
    """Set font for a run"""
    run.font.name = font_name
    run.font.size = Pt(size)
    # For Asian fonts compatibility
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)

def create_template():
    """Create Vietnamese thesis template"""
    doc = Document()

    # Set page margins (A4 format)
    sections = doc.sections
    for section in sections:
        section.page_width = Cm(21.0)   # A4 width
        section.page_height = Cm(29.7)  # A4 height
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)

    # Modify Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.space_after = Pt(6)

    # Modify Heading styles
    heading_sizes = {
        'Heading 1': 18,
        'Heading 2': 16,
        'Heading 3': 14,
        'Heading 4': 13,
    }

    for heading_name, size in heading_sizes.items():
        try:
            style = doc.styles[heading_name]
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(size)
            font.bold = True

            paragraph_format = style.paragraph_format
            paragraph_format.space_before = Pt(12)
            paragraph_format.space_after = Pt(6)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        except KeyError:
            pass

    # Modify Title style
    try:
        style = doc.styles['Title']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(18)
        font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except KeyError:
        pass

    # Modify TOC styles if available
    for i in range(1, 4):
        try:
            style = doc.styles[f'TOC {i}']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(13)
        except KeyError:
            pass

    # Add sample content to verify formatting
    title = doc.add_heading('Mẫu Tài Liệu Luận Văn', 0)

    doc.add_heading('Chương 1: Giới Thiệu', level=1)
    p = doc.add_paragraph('Đây là đoạn văn mẫu với font Times New Roman 13pt, dãn dòng 1.5.')

    doc.add_heading('1.1. Mục Tiêu', level=2)
    p = doc.add_paragraph('Nội dung mục tiêu nghiên cứu.')

    doc.add_heading('1.1.1. Chi Tiết', level=3)
    p = doc.add_paragraph('Nội dung chi tiết.')

    # Save template
    output_path = 'templates/reference.docx'
    doc.save(output_path)
    print(f"Template created: {output_path}")
    print("\nSpecifications:")
    print("- Font: Times New Roman")
    print("- Body: 13pt, Line spacing: 1.5")
    print("- Heading 1: 18pt bold")
    print("- Heading 2: 16pt bold")
    print("- Heading 3: 14pt bold")
    print("- Margins: Left 3.5cm, Right 2cm, Top/Bottom 3cm")
    print("- Page: A4 (210 × 297 mm)")

if __name__ == '__main__':
    create_template()
